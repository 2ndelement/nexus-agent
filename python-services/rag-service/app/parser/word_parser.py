"""
app/parser/word_parser.py — Word 解析器
"""
import logging
from pathlib import Path

from app.parser.base import BaseParser, ParseResult

logger = logging.getLogger(__name__)


class WordParser(BaseParser):
    """Word 文档解析器"""
    
    def supports(self, file_path: str) -> bool:
        """支持 Word 文件"""
        return Path(file_path).suffix.lower() in ['.docx', '.doc']
    
    def parse(self, file_path: str) -> ParseResult:
        """解析 Word 文档"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_parts = []
            
            # 提取标题
            title = ""
            if doc.core_properties.title:
                title = doc.core_properties.title
            else:
                title = Path(file_path).stem
            
            metadata = {
                "author": doc.core_properties.author or "",
                "source_file": file_path,
                "file_type": "word"
            }
            
            # 遍历所有元素
            for para in doc.paragraphs:
                if para.text.strip():
                    # 检测是否是标题
                    if para.style.name.startswith('Heading'):
                        text_parts.append(f"\n## {para.text}\n")
                    else:
                        text_parts.append(para.text)
            
            # 提取表格
            for table in doc.tables:
                text_parts.append("\n表格:")
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    text_parts.append(row_text)
            
            content = "\n".join(text_parts)
            
            return ParseResult(
                content=content,
                title=title,
                metadata=metadata,
                success=True
            )
            
        except ImportError:
            return ParseResult(
                content="",
                title=Path(file_path).stem,
                metadata={"file_type": "word"},
                success=False,
                error_message="python-docx not installed: pip install python-docx"
            )
        except Exception as e:
            logger.error(f"Word parsing error: {e}")
            return ParseResult(
                content="",
                title=Path(file_path).stem,
                metadata={"file_type": "word"},
                success=False,
                error_message=str(e)
            )
